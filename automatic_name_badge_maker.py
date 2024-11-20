"""To run this script in python:
1. Create the csv file of attendees and call this 'attendee_list.csv'. Add in any info to the
   'additional' column, e.g 'Plenary' or 'Committee'. Then check this csv carefully as some
   bookings appear twice (could be an error in the ZSL system or they double booked?)
2. Place in one folder: i) the RCUK logo and call this 'logo.PNG',  the 'attendee_list.csv', 
   and this python script.
3. Navigate to this folder
4. Run with the command:
   python automatic_name_badge_maker.py attendee_list.csv logo.PNG badges.docx
5. A word doc will be created with the name badges. Check this carefully."""


import csv
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import argparse

def process_csv(input_csv):
    attendees = []
    with open(input_csv, newline='', encoding='utf-8-sig') as csvfile:
        reader = csv.DictReader(csvfile, delimiter=',')
        for row in reader:
            # Extract and split the fields
            first_names = [name.strip() for name in row['First name of attendee'].split(';')]
            last_names = [name.strip() for name in row['Last name of attendee'].split(';')]
            organisations = [org.strip() for org in row['Organisation'].split(';')]

            # Handle the 'additional' column
            if 'additional' in row:
                additionals = [add.strip() for add in row['additional'].split(';')]
            else:
                additionals = [''] * len(first_names)  # Default to empty strings if column is missing

            max_len = max(len(first_names), len(last_names), len(organisations), len(additionals))

            # Extend lists to the same length
            first_names.extend([''] * (max_len - len(first_names)))
            last_names.extend([''] * (max_len - len(last_names)))
            organisations.extend([''] * (max_len - len(organisations)))
            additionals.extend([''] * (max_len - len(additionals)))

            for fn, ln, org, add in zip(first_names, last_names, organisations, additionals):
                attendee = {
                    'first_name': fn,
                    'last_name': ln,
                    'organisation': org,
                    'additional': add
                }
                attendees.append(attendee)
    return attendees

def create_name_badges(attendees, logo_path, output_docx):
    document = Document()
    sections = document.sections

    # Set page margins
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    # Badges layout configuration
    badges_per_row = 2
    badges_per_column = 4
    badges_per_page = badges_per_row * badges_per_column

    badge_width = Cm(9)
    badge_height = Cm(6)

    for i, attendee in enumerate(attendees):
        if i % badges_per_page == 0:
            if i > 0:
                document.add_page_break()
            table = document.add_table(rows=badges_per_column, cols=badges_per_row)
            table.autofit = False
            for row in table.rows:
                for cell in row.cells:
                    cell.width = badge_width
                    cell.height = badge_height  # Optionally set the height

        row_idx = (i % badges_per_page) // badges_per_row
        col_idx = (i % badges_per_page) % badges_per_row

        cell = table.cell(row_idx, col_idx)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        # Create a nested table within the cell
        inner_table = cell.add_table(rows=1, cols=2)
        inner_table.autofit = False

        # Set widths for the inner table columns
        inner_table.columns[0].width = Cm(4)  # Logo column
        inner_table.columns[1].width = Cm(5)  # Text column

        # Left cell (Logo)
        logo_cell = inner_table.cell(0, 0)
        logo_paragraph = logo_cell.paragraphs[0]
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, width=Cm(4))
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the logo if desired

        # Right cell (Name, Organisation, and Additional)
        text_cell = inner_table.cell(0, 1)
        text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Add attendee's name
        name_paragraph = text_cell.paragraphs[0]
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_paragraph.add_run(f"{attendee['first_name']} {attendee['last_name']}")
        name_run.font.size = Pt(18)
        name_run.bold = True

        # Add organisation
        org_paragraph = text_cell.add_paragraph()
        org_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        org_run = org_paragraph.add_run(attendee['organisation'])
        org_run.font.size = Pt(14)
        org_run.bold = True

        # Add additional info if available
        if attendee['additional']:
            add_paragraph = text_cell.add_paragraph()
            add_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run = add_paragraph.add_run(attendee['additional'])
            add_run.font.size = Pt(14)
            add_run.bold = True
            # Set font color to medium green (RGB: 0, 128, 0)
            add_run.font.color.rgb = RGBColor(0, 128, 0)

    document.save(output_docx)

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Generate name badges from a CSV file.')
    parser.add_argument('input_csv', help='Path to the input CSV file.')
    parser.add_argument('logo_path', help='Path to the logo image file.')
    parser.add_argument('output_docx', help='Path to the output DOCX file.')
    args = parser.parse_args()

    attendees = process_csv(args.input_csv)
    create_name_badges(attendees, args.logo_path, args.output_docx)
